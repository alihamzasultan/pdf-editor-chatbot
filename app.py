import os
import json
from flask import Flask, render_template, request, jsonify, session
from flask_cors import CORS
from dotenv import load_dotenv
import openai

# Load environment variables
load_dotenv()

# Initialize Flask app
app = Flask(__name__)
CORS(app)
app.secret_key = os.getenv("FLASK_SECRET_KEY", "default-secret-key")

# Initialize OpenAI client
client = openai.OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# Load questions from JSON
with open("faq.json", "r", encoding="utf-8") as f:
    question_data = json.load(f)["questions"]

# Helper to get the current question index
def get_current_question_index():
    return session.get("question_index", 0)

# Helper to get the next question
def get_next_question():
    index = get_current_question_index()
    if index < len(question_data):
        key = list(question_data[index].keys())[0]
        return question_data[index][key]
    return None

@app.route("/", methods=["GET"])
def index():
    session.clear()
    return render_template("index.html")

@app.route("/chatbot", methods=["POST"])
def chatbot():
    user_input = request.form.get("user_input", "").strip()
    if not user_input:
        return jsonify({"response": "Please enter a message."})

    if "question_index" not in session:
        session["question_index"] = 0
        session["conversation"] = []
        session["answers"] = {}

    current_index = session["question_index"]

    if current_index < len(question_data):
        question_entry = question_data[current_index]

        # Get the question key (e.g., "q2")
        question_key = [k for k in question_entry.keys() if not k.endswith("_type")][0]

        current_question = question_entry[question_key]

        # Extract type using the corresponding "{question_key}_type"
        question_type_key = f"{question_key}_type"
        question_type = question_entry.get(question_type_key, f"type_{current_index}")

        session["conversation"].append({
            "role": "user",
            "content": f"User was asked: '{current_question}'\nUser answered: '{user_input}'\nIs this a valid and complete answer? Reply only Yes or No."
        })

        try:
            completion = client.chat.completions.create(
                model="gpt-4o",
                messages=[{
                    "role": "system",
                    "content": "You are a helpful assistant that validates user answers to questions. You must reply with only one word: Yes or No. Do not include any explanation, don't be too strict to validate"

                }] + session["conversation"]
            )

            ai_response = completion.choices[0].message.content.strip().lower()
            print("GPT Response:", ai_response)


            # Clear the GPT check to keep conversation short
            session["conversation"] = []

            if ai_response.strip().lower() == "yes":

                # Save answer using type
                session["answers"][question_type] = user_input

                session["question_index"] += 1
                next_question = get_next_question()

                # If all questions answered, write to file
                if not next_question:
                    with open("answers.json", "w", encoding="utf-8") as f:
                        json.dump(session["answers"], f, indent=2)

                    fill_word_template()

                    session.pop("question_index", None)
                    session.pop("answers", None)
                    return jsonify({"response": "✅ All questions have been answered. Thank you!"})
                else:
                    return jsonify({"response": f"✅ Got it. Now, {next_question}"})

            else:
                return jsonify({"response": f"❌ Sorry, your answer wasn’t clear. Please try again: {current_question}"})

        except Exception as e:
            return jsonify({"response": f"Error: {str(e)}"})

    else:
        return jsonify({"response": "✅ All questions already answered. You may reset to start again."})

from docx import Document
import shutil

from docx import Document
import os
import shutil
import json
import threading
@app.route('/generate-doc', methods=['POST'])
def generate_doc():
    try:
        answers_path = os.path.join(os.getcwd(), 'answers.json')
        template_path = os.path.join(os.getcwd(), 'WFNJ-1JEnglish.docx')
        output_path = os.path.join(os.getcwd(), 'filled_template.docx')

        # Check if files exist
        if not os.path.exists(answers_path):
            return jsonify({'error': 'answers.json not found'}), 404
        if not os.path.exists(template_path):
            return jsonify({'error': 'Template Word file not found'}), 404

        # Load answers
        with open(answers_path, 'r', encoding='utf-8') as f:
            answers = json.load(f)
        print("Loaded answers:", answers)

        # Backup original
        backup_path = template_path.replace('.docx', '_backup.docx')
        shutil.copy(template_path, backup_path)
        print("Template backed up to:", backup_path)

        # Load and modify the document
        doc = Document(template_path)
        for para in doc.paragraphs:
            for key, value in answers.items():
                if f'{{{{{key}}}}}' in para.text:
                    para.text = para.text.replace(f'{{{{{key}}}}}', str(value))

        # Save the filled document
        doc.save(output_path)
        print("Filled document saved to:", output_path)
        timer = threading.Timer(300.0, delete_file, args=[output_path])
        timer.start()

        return jsonify({'message': 'Document generated successfully.'}), 200

    except Exception as e:
        print("Error:", str(e))
        return jsonify({'error': str(e)}), 500
def fill_word_template(answers_path='answers.json', template_path='WFNJ-1JEnglish.docx', output_path='filled_template.docx'):
    from docx import Document

    # Load answers
    with open(answers_path, 'r', encoding='utf-8') as f:
        answers = json.load(f)

    # Make a backup of the template
    backup_path = template_path.replace('.docx', '_backup.docx')
    shutil.copy(template_path, backup_path)

    # Load document
    doc = Document(template_path)

    # Replace placeholders
    for para in doc.paragraphs:
        for key, value in answers.items():
            if f'{{{{{key}}}}}' in para.text:
                para.text = para.text.replace(f'{{{{{key}}}}}', str(value))

    # Save filled document
    doc.save(output_path)
    print("Document filled and saved to:", output_path)

from flask import send_from_directory
from werkzeug.utils import safe_join
import os

@app.route('/download', methods=['GET'])
def download_file():
    directory = os.path.abspath(os.path.dirname(__file__))
    filename = 'filled_template.docx'
    file_path = safe_join(directory, filename)
    
    if not os.path.exists(file_path):
        return "File not found", 404
    
    # Schedule file deletion after 5 seconds
    timer = threading.Timer(5.0, delete_file, args=[file_path])
    timer.start()
    
    return send_from_directory(
        directory=directory,
        path=filename,
        as_attachment=True,
        download_name='GetMyAid_Results.docx'
    )

def delete_file(file_path):
    try:
        if os.path.exists(file_path):
            os.remove(file_path)
            print(f"Deleted file: {file_path}")
    except Exception as e:
        print(f"Error deleting file: {e}")

@app.route("/reset", methods=["POST"])
def reset_conversation():
    session.pop("conversation", None)
    session.pop("question_index", None)
    return jsonify({"message": "Conversation reset."})

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=int(os.environ.get("PORT", 5000)))
